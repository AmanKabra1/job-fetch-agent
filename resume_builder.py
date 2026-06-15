"""
Tailored resume builder for Aman Kabra.
========================================

Generate a job-specific resume in your own format. Paste a job description and
the builder auto-emphasizes the skills that match the role (reorders them to
the front of each category, bolds them, and rewrites the summary line to lead
with your top matching stack). Output is written to ./resume/ as PDF and/or
Word (.docx).

USAGE
-----
  # Generate from a job description file (recommended):
  python resume_builder.py new --title "Backend Developer" --company "Acme" --jd jd.txt

  # Or paste/point the JD inline:
  python resume_builder.py new --title "Node.js Developer" --company "Acme" \
      --jd-text "We need a Node.js / NestJS engineer with MySQL and AWS ..."

  # No JD -> plain base resume:
  python resume_builder.py new --title "Software Engineer" --company "Acme"

  # Pick format(s): pdf | docx | both (default: both)
  python resume_builder.py new --title "..." --company "..." --jd jd.txt --format pdf

  # List everything you've generated:
  python resume_builder.py list

  # Delete a generated resume (both formats for that base name), or everything:
  python resume_builder.py delete "Aman_Kabra_Acme_Backend-Developer_2026-06-15"
  python resume_builder.py delete --all

Edit your details in resume_profile.py — both formats stay in sync.
"""

import os
import re
import sys
import argparse
import datetime as dt

import resume_profile as P

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume")


# --------------------------------------------------------------------------- #
# TAILORING
# --------------------------------------------------------------------------- #
def find_matched_skills(jd_text: str) -> set:
    """Return the set of SKILLS labels whose aliases appear in the JD text."""
    if not jd_text:
        return set()
    text = " " + jd_text.lower() + " "
    matched = set()
    for label, aliases in P.SKILL_KEYWORDS.items():
        for alias in aliases:
            # word-boundary match so "go" doesn't fire on "google", "rest" not "restaurant"
            pattern = r"(?<![a-z0-9])" + re.escape(alias.strip()) + r"(?![a-z0-9])"
            if re.search(pattern, text):
                matched.add(label)
                break
    return matched


def tailor_skills(matched: set):
    """Reorder each skill category so matched skills come first (stable)."""
    tailored = {}
    for category, items in P.SKILLS.items():
        hits = [s for s in items if s in matched]
        rest = [s for s in items if s not in matched]
        tailored[category] = hits + rest
    return tailored


def build_summary(matched: set) -> str:
    """Rewrite the summary line to lead with the top matching stack."""
    # Preferred display order for the summary lead-in.
    priority = [
        "NestJS", "Node.js", "TypeScript", "Express.js", "Python",
        "Java (Spring Boot)", "Spring Boot", "FastAPI", "Django", "Go (Golang)",
        "Angular 17 (Signals, standalone components, RxJS)",
        "Microservices", "RESTful APIs", "AWS S3", "Docker",
        "LLM Integration", "RAG (Retrieval-Augmented Generation)",
    ]
    # Friendly short names for the summary sentence.
    short = {
        "Java (Spring Boot)": "Java",
        "Go (Golang)": "Go",
        "Angular 17 (Signals, standalone components, RxJS)": "Angular",
        "RAG (Retrieval-Augmented Generation)": "RAG",
        "AWS S3": "AWS",
    }
    lead = [s for s in priority if s in matched]
    if len(lead) < 3:
        # Top up with defaults so the sentence always reads well.
        for s in P.DEFAULT_STACK:
            if s not in lead:
                lead.append(s)
    lead = [short.get(s, s) for s in lead][:5]
    stack = ", ".join(lead[:-1]) + ", and " + lead[-1] if len(lead) > 1 else lead[0]
    return P.SUMMARY_TEMPLATE.format(stack=stack)


def slugify(text: str) -> str:
    text = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-")
    return text or "x"


def base_filename(company: str, title: str) -> str:
    today = dt.date.today().isoformat()
    return f"{slugify(P.NAME)}_{slugify(company)}_{slugify(title)}_{today}"


# --------------------------------------------------------------------------- #
# PDF RENDERER  (reportlab)
# --------------------------------------------------------------------------- #
def render_pdf(path, summary, skills, matched, target_title, target_company):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
        ListFlowable, ListItem,
    )
    from reportlab.lib import colors

    DARK = colors.HexColor("#1A1A1A")
    ACCENT = colors.HexColor("#1E3A5F")

    doc = SimpleDocTemplate(
        path, pagesize=A4,
        leftMargin=15 * mm, rightMargin=15 * mm,
        topMargin=12 * mm, bottomMargin=12 * mm,
        title=f"{P.NAME} - Resume", author=P.NAME,
    )
    content_w = doc.width

    name_st = ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=20,
                             alignment=TA_CENTER, textColor=DARK, spaceAfter=2, leading=23)
    contact_st = ParagraphStyle("contact", fontName="Helvetica", fontSize=8.5,
                                alignment=TA_CENTER, textColor=DARK, leading=11)
    summary_st = ParagraphStyle("summary", fontName="Helvetica", fontSize=9.2,
                                alignment=TA_JUSTIFY, textColor=DARK, leading=12.5)
    body_st = ParagraphStyle("body", fontName="Helvetica", fontSize=9.2,
                             textColor=DARK, leading=12.5)
    bullet_st = ParagraphStyle("bullet", parent=body_st, leftIndent=10, bulletIndent=0)
    role_l = ParagraphStyle("role_l", fontName="Helvetica-Bold", fontSize=10, textColor=DARK)
    role_r = ParagraphStyle("role_r", fontName="Helvetica-Bold", fontSize=9.2,
                            textColor=DARK, alignment=2)
    sub_l = ParagraphStyle("sub_l", fontName="Helvetica-Oblique", fontSize=9, textColor=DARK)
    sub_r = ParagraphStyle("sub_r", fontName="Helvetica-Oblique", fontSize=9,
                           textColor=DARK, alignment=2)

    story = []

    def section(title):
        story.append(Spacer(1, 7))
        story.append(Paragraph(title.upper(),
                     ParagraphStyle("sec", fontName="Helvetica-Bold", fontSize=10.5,
                                    textColor=ACCENT, spaceAfter=2, leading=12)))
        story.append(HRFlowable(width="100%", thickness=0.8, color=ACCENT,
                                spaceBefore=1, spaceAfter=4))

    def two_col(left, right, lstyle, rstyle):
        t = Table([[Paragraph(left, lstyle), Paragraph(right, rstyle)]],
                  colWidths=[content_w * 0.68, content_w * 0.32])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        return t

    def bullets(items):
        return ListFlowable(
            [ListItem(Paragraph(b, bullet_st), value="bullet", leftIndent=12)
             for b in items],
            bulletType="bullet", bulletFontSize=7, start="•",
            leftIndent=12, spaceBefore=1, spaceAfter=1,
        )

    def bold_matched(skill):
        """Bold a skill label in PDF markup if it matched the JD."""
        return f"<b>{skill}</b>" if skill in matched else skill

    # Header
    story.append(Paragraph(P.NAME, name_st))
    contact = " &nbsp;|&nbsp; ".join([
        P.LOCATION, P.PHONE,
        f'<a href="mailto:{P.EMAIL}">{P.EMAIL}</a>',
        f'<a href="{P.LINKEDIN}">LinkedIn</a>',
        f'<a href="{P.GITHUB}">GitHub</a>',
    ])
    story.append(Paragraph(contact, contact_st))

    # Professional summary
    section("Professional Summary")
    story.append(Paragraph(summary, summary_st))

    # Experience
    section("Experience")
    for i, job in enumerate(P.EXPERIENCE):
        if i:
            story.append(Spacer(1, 4))
        story.append(two_col(job["company"], job["dates"], role_l, role_r))
        story.append(two_col(job["title"], job["location"], sub_l, sub_r))
        story.append(Spacer(1, 1))
        story.append(bullets(job["bullets"]))

    # Technical skills
    section("Technical Skills")
    for category, items in skills.items():
        line = f"<b>{category}:</b> " + ", ".join(bold_matched(s) for s in items)
        story.append(Paragraph(line, body_st))
        story.append(Spacer(1, 1.5))

    # Projects
    section("Projects")
    for i, proj in enumerate(P.PROJECTS):
        if i:
            story.append(Spacer(1, 4))
        head = (f'<b>{proj["name"]}</b> | <i>{proj["stack"]}</i> | '
                f'<a href="{proj["link"]}">{proj["link"]}</a>')
        story.append(Paragraph(head, body_st))
        story.append(Spacer(1, 1))
        story.append(bullets(proj["bullets"]))

    # Education
    section("Education")
    story.append(two_col(P.EDUCATION["school"], P.EDUCATION["dates"], role_l, role_r))
    story.append(two_col(P.EDUCATION["degree"], P.EDUCATION["location"], sub_l, sub_r))

    # Certifications
    section("Certifications")
    story.append(Paragraph("  |  ".join(P.CERTIFICATIONS), body_st))

    doc.build(story)


# --------------------------------------------------------------------------- #
# WORD RENDERER  (python-docx)
# --------------------------------------------------------------------------- #
def render_docx(path, summary, skills, matched, target_title, target_company):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    DARK = RGBColor(0x1A, 0x1A, 0x1A)
    ACCENT = RGBColor(0x1E, 0x3A, 0x5F)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    usable_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = DARK

    def no_space(p, before=0, after=2):
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.0
        return p

    def add_bottom_border(p):
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1E3A5F")
        pbdr.append(bottom)
        pPr.append(pbdr)

    def section_heading(title):
        p = no_space(doc.add_paragraph(), before=6, after=3)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = ACCENT
        add_bottom_border(p)

    def two_col(left, right, left_bold=True, italic=False):
        p = no_space(doc.add_paragraph())
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
        lr = p.add_run(left)
        lr.bold = left_bold
        lr.italic = italic
        p.add_run("\t")
        rr = p.add_run(right)
        rr.italic = italic
        if not italic:
            rr.bold = left_bold
        return p

    def bullet(text):
        p = no_space(doc.add_paragraph(style="List Bullet"), after=1)
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(text)
        return p

    # Header
    p = no_space(doc.add_paragraph(), after=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(P.NAME)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = DARK

    p = no_space(doc.add_paragraph(), after=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = p.add_run(" | ".join([P.LOCATION, P.PHONE, P.EMAIL, P.LINKEDIN, P.GITHUB]))
    c.font.size = Pt(8)

    # Summary
    section_heading("Professional Summary")
    no_space(doc.add_paragraph()).add_run(summary)

    # Experience
    section_heading("Experience")
    for job in P.EXPERIENCE:
        two_col(job["company"], job["dates"], left_bold=True, italic=False)
        two_col(job["title"], job["location"], left_bold=False, italic=True)
        for b in job["bullets"]:
            bullet(b)

    # Technical skills
    section_heading("Technical Skills")
    for category, items in skills.items():
        p = no_space(doc.add_paragraph(), after=2)
        cr = p.add_run(f"{category}: ")
        cr.bold = True
        for idx, s in enumerate(items):
            run = p.add_run(s + (", " if idx < len(items) - 1 else ""))
            if s in matched:
                run.bold = True

    # Projects
    section_heading("Projects")
    for proj in P.PROJECTS:
        p = no_space(doc.add_paragraph())
        nr = p.add_run(proj["name"])
        nr.bold = True
        sr = p.add_run(f' | {proj["stack"]} | ')
        sr.italic = True
        p.add_run(proj["link"]).font.color.rgb = ACCENT
        for b in proj["bullets"]:
            bullet(b)

    # Education
    section_heading("Education")
    two_col(P.EDUCATION["school"], P.EDUCATION["dates"], left_bold=True, italic=False)
    two_col(P.EDUCATION["degree"], P.EDUCATION["location"], left_bold=False, italic=True)

    # Certifications
    section_heading("Certifications")
    no_space(doc.add_paragraph()).add_run("  |  ".join(P.CERTIFICATIONS))

    doc.save(path)


# --------------------------------------------------------------------------- #
# COMMANDS
# --------------------------------------------------------------------------- #
def read_jd(args) -> str:
    if args.jd_text:
        return args.jd_text
    if args.jd:
        if not os.path.exists(args.jd):
            sys.exit(f"JD file not found: {args.jd}")
        with open(args.jd, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    return ""


def cmd_new(args):
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    jd_text = read_jd(args)
    matched = find_matched_skills(jd_text)
    skills = tailor_skills(matched)
    summary = build_summary(matched)

    base = base_filename(args.company, args.title)
    fmt = args.format
    written = []

    if fmt in ("pdf", "both"):
        pdf_path = os.path.join(OUTPUT_DIR, base + ".pdf")
        render_pdf(pdf_path, summary, skills, matched, args.title, args.company)
        written.append(pdf_path)
    if fmt in ("docx", "both"):
        docx_path = os.path.join(OUTPUT_DIR, base + ".docx")
        render_docx(docx_path, summary, skills, matched, args.title, args.company)
        written.append(docx_path)

    print(f"Target role : {args.title} @ {args.company}")
    if jd_text:
        if matched:
            print(f"Emphasized  : {', '.join(sorted(matched))}")
        else:
            print("Emphasized  : (no skill keywords matched in the JD)")
    else:
        print("Emphasized  : (no JD supplied - base resume)")
    print("Created:")
    for w in written:
        print(f"  - {w}")


def cmd_list(args):
    if not os.path.isdir(OUTPUT_DIR):
        print("No resumes generated yet.")
        return
    files = sorted(f for f in os.listdir(OUTPUT_DIR) if f.endswith((".pdf", ".docx")))
    if not files:
        print("No resumes generated yet.")
        return
    print(f"Resumes in {OUTPUT_DIR}:")
    for f in files:
        full = os.path.join(OUTPUT_DIR, f)
        size = os.path.getsize(full) / 1024
        print(f"  - {f}  ({size:.0f} KB)")


def cmd_delete(args):
    if not os.path.isdir(OUTPUT_DIR):
        print("Nothing to delete.")
        return

    if args.all:
        targets = [f for f in os.listdir(OUTPUT_DIR) if f.endswith((".pdf", ".docx"))]
    elif args.name:
        # Accept a base name (delete both formats) or an exact filename.
        name = args.name
        targets = []
        for f in os.listdir(OUTPUT_DIR):
            if f == name or os.path.splitext(f)[0] == os.path.splitext(name)[0]:
                targets.append(f)
    else:
        sys.exit("Provide a resume name to delete, or use --all.")

    if not targets:
        print(f"No matching resume found for: {args.name}")
        return

    for f in targets:
        os.remove(os.path.join(OUTPUT_DIR, f))
        print(f"Deleted: {f}")


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def build_parser():
    parser = argparse.ArgumentParser(
        description="Build a tailored resume for a specific job (PDF + Word).",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)

    p_new = sub.add_parser("new", help="generate a tailored resume")
    p_new.add_argument("--title", required=True, help="target job title")
    p_new.add_argument("--company", required=True, help="target company")
    p_new.add_argument("--jd", help="path to a job description text file")
    p_new.add_argument("--jd-text", dest="jd_text", help="job description text inline")
    p_new.add_argument("--format", choices=["pdf", "docx", "both"], default="both",
                       help="output format (default: both)")
    p_new.set_defaults(func=cmd_new)

    p_list = sub.add_parser("list", help="list generated resumes")
    p_list.set_defaults(func=cmd_list)

    p_del = sub.add_parser("delete", help="delete a generated resume (or --all)")
    p_del.add_argument("name", nargs="?", help="resume base name or filename to delete")
    p_del.add_argument("--all", action="store_true", help="delete every generated resume")
    p_del.set_defaults(func=cmd_delete)

    return parser


def main():
    args = build_parser().parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
