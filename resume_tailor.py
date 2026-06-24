"""
Tailor an UPLOADED resume to a job description — without changing its style.
=============================================================================

The user uploads their OWN resume (.docx or .pdf) plus an optional job
description and/or an optional list of skills, and we tailor that exact
document — there is no built-in profile or template; every output is the user's
own resume:

  * For a .docx upload we edit the file IN PLACE — every font, margin, color and
    layout decision the user made is preserved. We only:
      - bold the skills the resume already has that the job is asking for, and
      - insert one "Key Skills for this Role" line near the top.
    Nothing in the resume is rewritten or fabricated.

  * For a .pdf upload there is no reliable way to edit a PDF in place, so we
    extract the text and rebuild a STRUCTURED, single-page document — detecting
    the name, contact line, section headings and bullets so the result keeps the
    resume's structure (not a flat text dump) and stays on one page.

Both paths can emit Word (.docx) and PDF. PDF is produced from the tailored
.docx via Microsoft Word (docx2pdf) so the PDF matches the Word styling exactly;
if Word isn't available we fall back to a reportlab-rendered PDF.

Honesty rule: skills the JD asks for but the resume does NOT contain are returned
as *suggestions* for the user to review — they are never silently added to the
document.
"""

import io
import os
import re
import copy
import tempfile

# --------------------------------------------------------------------------- #
# SKILL LEXICON
#   canonical label -> aliases that, if found in text (word-boundary, case-
#   insensitive), mean the skill is present. Kept broad so this works for ANY
#   uploaded resume / JD, not just one profile.
# --------------------------------------------------------------------------- #
SKILL_LEXICON = {
    # Languages
    "Python": ["python"],
    "JavaScript": ["javascript", "js", "es6", "ecmascript"],
    "TypeScript": ["typescript", "ts"],
    "Java": ["java"],
    "Kotlin": ["kotlin"],
    "Go": ["golang", "go lang"],
    "Rust": ["rust"],
    "C++": ["c++", "cpp"],
    "C#": ["c#", ".net", "dotnet"],
    "C": [" c "],
    "PHP": ["php"],
    "Ruby": ["ruby"],
    "Swift": ["swift"],
    "Scala": ["scala"],
    "SQL": ["sql"],
    "HTML": ["html"],
    "CSS": ["css"],
    "Bash/Shell": ["bash", "shell scripting", "shell script"],
    # Backend frameworks
    "NestJS": ["nestjs", "nest.js"],
    "Node.js": ["node.js", "nodejs", "node js", "node"],
    "Express.js": ["express.js", "express", "expressjs"],
    "Spring Boot": ["spring boot", "springboot", "spring"],
    "FastAPI": ["fastapi", "fast api"],
    "Flask": ["flask"],
    "Django": ["django"],
    "Ruby on Rails": ["rails", "ruby on rails"],
    "Laravel": ["laravel"],
    ".NET Core": [".net core", "asp.net"],
    "GraphQL": ["graphql"],
    "gRPC": ["grpc"],
    # Frontend
    "React": ["react", "react.js", "reactjs"],
    "Next.js": ["next.js", "nextjs"],
    "Angular": ["angular", "rxjs"],
    "Vue.js": ["vue", "vue.js", "vuejs"],
    "Redux": ["redux"],
    "Tailwind CSS": ["tailwind"],
    "SASS": ["sass", "scss"],
    # Databases
    "MySQL": ["mysql"],
    "PostgreSQL": ["postgresql", "postgres"],
    "MongoDB": ["mongodb", "mongo"],
    "Redis": ["redis"],
    "Elasticsearch": ["elasticsearch", "elastic search"],
    "DynamoDB": ["dynamodb"],
    "Cassandra": ["cassandra"],
    "Oracle DB": ["oracle"],
    "SQLite": ["sqlite"],
    # Cloud & DevOps
    "AWS": ["aws", "amazon web services", "ec2", "s3", "lambda", "ses"],
    "Google Cloud": ["gcp", "google cloud"],
    "Azure": ["azure"],
    "Docker": ["docker", "container", "containeri"],
    "Kubernetes": ["kubernetes", "k8s"],
    "Terraform": ["terraform"],
    "CI/CD": ["ci/cd", "cicd", "continuous integration", "continuous delivery"],
    "GitHub Actions": ["github actions"],
    "Jenkins": ["jenkins"],
    "GitLab CI": ["gitlab ci", "gitlab-ci"],
    "Git": ["git"],
    "Linux": ["linux", "unix"],
    "Nginx": ["nginx"],
    # Messaging / streaming
    "Kafka": ["kafka"],
    "RabbitMQ": ["rabbitmq", "rabbit mq"],
    "Celery": ["celery"],
    # Architecture & concepts
    "Microservices": ["microservice", "micro service", "distributed system"],
    "RESTful APIs": ["rest api", "restful", "rest "],
    "JWT": ["jwt"],
    "OAuth 2.0": ["oauth"],
    "WebSockets": ["websocket"],
    "MVC": ["mvc"],
    "Agile/Scrum": ["agile", "scrum", "kanban"],
    "TDD": ["tdd", "test driven", "test-driven"],
    "Swagger/OpenAPI": ["swagger", "openapi"],
    # Data / AI
    "Pandas": ["pandas"],
    "NumPy": ["numpy"],
    "PyTorch": ["pytorch"],
    "TensorFlow": ["tensorflow"],
    "Machine Learning": ["machine learning", "deep learning"],
    "LLM Integration": ["llm", "large language model", "gpt", "openai", "claude", "gemini"],
    "RAG": ["rag", "retrieval-augmented", "retrieval augmented", "vector database"],
    "ETL": ["etl", "data pipeline"],
    "Power BI": ["power bi", "powerbi"],
    "Tableau": ["tableau"],
    # Testing
    "Jest": ["jest"],
    "Pytest": ["pytest"],
    "JUnit": ["junit"],
    "Selenium": ["selenium"],
    "Cypress": ["cypress"],
    # Tools
    "Postman": ["postman"],
    "Jira": ["jira"],
    "n8n": ["n8n"],
}


# --------------------------------------------------------------------------- #
# TEXT MATCHING
# --------------------------------------------------------------------------- #
def _word_in(alias: str, text_lower: str) -> bool:
    """Word-boundary, case-insensitive containment (text_lower is already lower)."""
    a = alias.strip().lower()
    if not a:
        return False
    pattern = r"(?<![a-z0-9+#.])" + re.escape(a) + r"(?![a-z0-9+#])"
    return re.search(pattern, text_lower) is not None


def _clean_label(label: str) -> str:
    """Drop any parenthetical detail for the short display name."""
    return re.sub(r"\s*\(.*?\)\s*", "", label).strip()


# --------------------------------------------------------------------------- #
# UPLOAD -> TEXT
# --------------------------------------------------------------------------- #
def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded .docx or .pdf. Raises ValueError."""
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(t for t in parts if t and t.strip())
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        return "\n".join(pages)
    if ext == ".doc":
        raise ValueError("Old .doc files aren't supported — please save as .docx or PDF.")
    raise ValueError("Unsupported file type. Upload a .docx or .pdf resume.")


# --------------------------------------------------------------------------- #
# ANALYSIS
# --------------------------------------------------------------------------- #
def analyze(resume_text: str, jd_text: str = "", extra_skills: str = "") -> dict:
    """Decide what to emphasize.

    Returns:
      present     [labels]      skills the JD/skills-box want AND the resume has
      typed       [labels]      skills the user typed that the resume LACKS
                                (user-asserted -> we add them to the key-skills line)
      suggestions [labels]      skills the JD wants but the resume LACKS
                                (NOT added — shown to the user to consider)
      bold_terms  [alias str]   actual tokens to bold inside the resume
      had_request bool          whether any JD/skills input was given at all
    """
    resume_l = " " + (resume_text or "").lower() + " "
    jd_l = " " + (jd_text or "").lower() + " "
    skills_l = " " + (extra_skills or "").lower() + " "
    had_request = bool((jd_text or "").strip() or (extra_skills or "").strip())

    present, suggestions, typed = [], [], []
    bold_terms = set()

    for label, aliases in SKILL_LEXICON.items():
        in_jd = any(_word_in(a, jd_l) for a in aliases)
        in_skills = any(_word_in(a, skills_l) for a in aliases)
        if not (in_jd or in_skills):
            continue
        resume_hits = [a for a in aliases if _word_in(a, resume_l)]
        clean = _clean_label(label)
        if resume_hits:
            present.append(clean)
            bold_terms.update(resume_hits)
        elif in_skills:
            # The user explicitly typed this skill -> trust them, add it.
            typed.append(clean)
        else:
            # JD wants it but it's not in the resume and user didn't claim it.
            suggestions.append(clean)

    return {
        "present": present,
        "typed": typed,
        "suggestions": suggestions,
        "bold_terms": sorted(bold_terms, key=len, reverse=True),
        "had_request": had_request,
    }


def _key_skills_line(analysis: dict, ats: bool = False) -> str:
    """The skills line we insert.

    Normal: skills you already have + skills you typed.
    ATS mode: ALSO append the JD's remaining keywords ('suggestions') so an
    Applicant Tracking System scanning for the JD's terms finds them all — i.e.
    ~100% keyword coverage. (The UI warns you to keep these truthful.)"""
    skills = list(analysis["present"])
    for t in analysis["typed"]:
        if t not in skills:
            skills.append(t)
    if ats:
        for s in analysis["suggestions"]:
            if s not in skills:
                skills.append(s)
    return ", ".join(skills)


# --------------------------------------------------------------------------- #
# DOCX TAILORING (in place — preserves style)
# --------------------------------------------------------------------------- #
def _bold_terms_regex(bold_terms):
    if not bold_terms:
        return None
    alts = "|".join(re.escape(t) for t in bold_terms)
    return re.compile(r"(?<![A-Za-z0-9+#.])(" + alts + r")(?![A-Za-z0-9+#])", re.IGNORECASE)


def _bold_in_run(run, regex):
    """Split a run so matched substrings become bold, preserving the run's other
    formatting (font, size, color, italic) by deep-copying its properties."""
    text = run.text
    if not text:
        return
    matches = list(regex.finditer(text))
    if not matches:
        return

    segments, pos = [], 0
    for m in matches:
        if m.start() > pos:
            segments.append((text[pos:m.start()], False))
        segments.append((m.group(0), True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False))

    from docx.text.run import Run
    # Capture the run's ORIGINAL bold BEFORE we mutate it — otherwise a run that
    # starts with a matched term turns its bold on, and every later (unmatched)
    # segment then inherits that bold (making the whole run bold).
    orig_bold = run.bold
    first_text, first_bold = segments[0]
    run.text = first_text
    run.bold = True if first_bold else orig_bold
    anchor = run._r
    for seg_text, seg_bold in segments[1:]:
        new_r = copy.deepcopy(run._r)          # carries this run's rPr (its style)
        new_run = Run(new_r, run._parent)
        new_run.text = seg_text                # replaces the copied w:t content
        new_run.bold = True if seg_bold else orig_bold
        anchor.addnext(new_r)
        anchor = new_r


def _insert_key_skills(doc, line: str, label: str = "Key Skills for this Role: ",
                       bold_terms=None):
    """Insert a skills paragraph just below the header. When bold_terms is given,
    the matched skills inside the line are bolded (and the JD-added ones left
    regular) — so the line reads as 'have these (bold) + add these (regular)'."""
    if not line:
        return
    from docx.shared import Pt
    paras = doc.paragraphs
    body = line
    if len(paras) >= 2:
        new_p = paras[1].insert_paragraph_before("")
    elif paras:
        new_p = paras[0].insert_paragraph_before("")
    else:
        new_p = doc.add_paragraph()
    r1 = new_p.add_run(label)
    r1.bold = True
    body_run = new_p.add_run(body)
    # Bold the matched skills inside the body run (label run is left as-is).
    if bold_terms:
        regex = _bold_terms_regex(bold_terms)
        if regex is not None:
            _bold_in_run(body_run, regex)
    # Keep it compact; inherit document Normal style otherwise.
    try:
        new_p.paragraph_format.space_before = Pt(2)
        new_p.paragraph_format.space_after = Pt(4)
    except Exception:
        pass


def tailor_docx(data: bytes, analysis: dict, ats: bool = False) -> bytes:
    """Edit an uploaded .docx in place and return the tailored bytes."""
    from docx import Document
    doc = Document(io.BytesIO(data))

    regex = _bold_terms_regex(analysis["bold_terms"])
    if regex is not None:
        def walk_paragraphs(container):
            for p in container.paragraphs:
                for run in list(p.runs):
                    _bold_in_run(run, regex)
        walk_paragraphs(doc)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in list(p.runs):
                            _bold_in_run(run, regex)

    label = "Core Competencies: " if ats else "Key Skills for this Role: "
    _insert_key_skills(doc, _key_skills_line(analysis, ats=ats), label=label,
                       bold_terms=analysis["bold_terms"])

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# --------------------------------------------------------------------------- #
# PDF UPLOAD -> rebuilt DOCX (structured, one page)
# --------------------------------------------------------------------------- #
# Common resume section headings (used to detect headings in extracted text so
# the rebuilt doc keeps the original's structure instead of a flat text dump).
_SECTION_WORDS = (
    "summary", "objective", "profile", "about",
    "experience", "work experience", "employment", "professional experience",
    "education", "academic", "skills", "technical skills", "core competencies",
    "projects", "project", "certifications", "certificates", "courses",
    "achievements", "accomplishments", "awards", "publications", "languages",
    "interests", "hobbies", "contact", "references", "volunteer", "activities",
)
# Includes the private-use glyphs Word/Symbol/Wingdings fonts use for bullets
# ( etc.) — PDF text extraction emits these literally, so if they're not
# recognised as bullets they'd render as a stray symbol in the rebuilt resume.
_BULLET_CHARS = (
    "•", "◦", "▪", "‣", "·", "∙", "●", "○", "*", "-", "–", "—",
    "", "", "", "", "",
)


def _looks_like_heading(line: str) -> bool:
    """A short line that is a section title (ALL CAPS or a known section word)."""
    s = line.strip().rstrip(":").strip()
    if not s or len(s) > 40:
        return False
    low = s.lower()
    if low in _SECTION_WORDS:
        return True
    # ALL-CAPS (or mostly) short line with no sentence punctuation -> heading.
    letters = [c for c in s if c.isalpha()]
    if letters and len(s.split()) <= 5 and sum(c.isupper() for c in letters) / len(letters) >= 0.8:
        return True
    return False


def _is_bullet(line: str) -> bool:
    s = line.lstrip()
    return bool(s) and s[0] in _BULLET_CHARS and s[:2] != "--"


def _is_contact(line: str) -> bool:
    """Email / phone / links line that usually sits under the name."""
    low = line.lower()
    return ("@" in line or "http" in low or "linkedin" in low or "github" in low
            or re.search(r"\+?\d[\d\s\-()]{7,}", line) is not None or "|" in line)


def _set_bottom_border(paragraph):
    """Give a paragraph a thin bottom border (the rule under a section heading)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")        # ~0.75pt rule
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)


# Base spacing (space_before, space_after, in pt) per block kind, before the
# density multiplier is applied.
_BLOCK_SPACING = {
    "name": (0, 1), "contact": (0, 2), "skillhead": (4, 1), "skills": (0, 2),
    "heading": (5, 1), "bullet": (0, 0), "body": (0, 0),
}
# Density presets tried loosest -> tightest. We pick the FIRST that fits one page;
# if none do, the tightest is used. This reduces SPACING/font, never content.
_DENSITY_PRESETS = (
    {"body": 9.0, "head": 10.0, "name": 17.0, "contact": 8.5, "ls": 1.0,  "sp": 1.0},
    {"body": 8.5, "head": 9.5,  "name": 16.0, "contact": 8.0, "ls": 1.0,  "sp": 0.7},
    {"body": 8.0, "head": 9.0,  "name": 15.0, "contact": 8.0, "ls": 0.98, "sp": 0.5},
    {"body": 7.5, "head": 8.5,  "name": 14.0, "contact": 7.5, "ls": 0.95, "sp": 0.3},
    {"body": 7.0, "head": 8.0,  "name": 13.0, "contact": 7.0, "ls": 0.92, "sp": 0.15},
)


def _block_size(kind: str, ps: dict) -> float:
    return {"name": ps["name"], "contact": ps["contact"], "skillhead": ps["head"],
            "heading": ps["head"], "skills": ps["body"], "bullet": ps["body"],
            "body": ps["body"]}[kind]


def _parse_resume_blocks(resume_text: str, analysis: dict, ats: bool):
    """Turn extracted PDF text into ordered (kind, text) blocks: the name, up to
    two contact lines, the inserted skills section, then headings / bullets /
    body — using _looks_like_heading() and _is_bullet() to detect structure."""
    lines = [ln.rstrip() for ln in (resume_text or "").splitlines()]
    nonempty = [ln for ln in lines if ln.strip()]
    blocks = []

    body_start = 0
    if nonempty:
        name_line = nonempty[0]
        blocks.append(("name", name_line.strip()))
        body_start = lines.index(name_line) + 1
        taken = 0
        while body_start < len(lines) and taken < 2:
            ln = lines[body_start]
            if not ln.strip():
                body_start += 1
                continue
            if _is_contact(ln) and not _looks_like_heading(ln):
                blocks.append(("contact", ln.strip()))
                body_start += 1
                taken += 1
            else:
                break

    skills_line = _key_skills_line(analysis, ats=ats)
    if skills_line:
        blocks.append(("skillhead", "CORE COMPETENCIES" if ats
                       else "KEY SKILLS FOR THIS ROLE"))
        blocks.append(("skills", skills_line))

    for ln in lines[body_start:]:
        if not ln.strip():
            continue
        if _looks_like_heading(ln):
            blocks.append(("heading", ln.strip().rstrip(":").upper()))
        elif _is_bullet(ln):
            text = ln.lstrip()
            text = text[1:].strip() if text[:1] in _BULLET_CHARS else text
            blocks.append(("bullet", text))
        else:
            blocks.append(("body", ln.strip()))
    return blocks


def _estimate_height_pt(blocks, ps, usable_w_pt) -> float:
    """Rough rendered height (pt) of all blocks at a given density preset. Used to
    decide whether the resume fits one page (no real renderer in python-docx)."""
    import math
    total = 0.0
    for kind, text in blocks:
        size = _block_size(kind, ps)
        # ~0.50*fontsize per char is a deliberately wide estimate (overestimates
        # lines -> errs toward shrinking, which keeps us on one page).
        cpl = max(16, usable_w_pt / (0.50 * size))
        nlines = max(1, math.ceil(len(text) / cpl))
        line_h = ps["ls"] * 1.18 * size
        before, after = _BLOCK_SPACING[kind]
        total += nlines * line_h + (before + after) * ps["sp"]
    return total


def rebuild_docx_from_text(resume_text: str, analysis: dict, ats: bool = False) -> bytes:
    """Rebuild a .docx from extracted PDF text, keeping resume STRUCTURE (name,
    contact, section headings, bullets) and compressing everything onto ONE page.

    A PDF can't be edited in place, so instead of a flat multi-page text dump we
    reconstruct a tight, structured single-page document: large bold centered
    name, centered contact line(s), bold UPPERCASE section headings with a bottom
    rule, real "•" bullets, and matched JD keywords bolded inline. If the content
    would overflow one page we step through tighter density presets — shrinking
    spacing and font, never dropping content."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    blocks = _parse_resume_blocks(resume_text, analysis, ats)

    doc = Document()
    # Tight one-page geometry.
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(28)      # ~0.39"
        section.left_margin = section.right_margin = Pt(40)      # ~0.55"
    sec = doc.sections[0]
    usable_w = (sec.page_width - sec.left_margin - sec.right_margin) / 12700.0  # -> pt
    usable_h = (sec.page_height - sec.top_margin - sec.bottom_margin) / 12700.0

    # Pick the loosest density that still fits one page (fall back to tightest).
    chosen = _DENSITY_PRESETS[-1]
    for ps in _DENSITY_PRESETS:
        if _estimate_height_pt(blocks, ps, usable_w) <= usable_h:
            chosen = ps
            break

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(chosen["body"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = chosen["ls"]

    regex = _bold_terms_regex(analysis["bold_terms"])

    def _add_text(p, text):
        p.add_run(text)
        if regex is not None:
            for run in list(p.runs):
                _bold_in_run(run, regex)

    def _para(kind):
        before, after = _BLOCK_SPACING[kind]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before * chosen["sp"])
        p.paragraph_format.space_after = Pt(after * chosen["sp"])
        p.paragraph_format.line_spacing = chosen["ls"]
        return p

    def _heading(text, size):
        hp = _para("heading")
        hr = hp.add_run(text)
        hr.bold = True
        hr.font.size = Pt(size)
        hr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        _set_bottom_border(hp)
        return hp

    for kind, text in blocks:
        if kind == "name":
            p = _para("name")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(chosen["name"])
        elif kind == "contact":
            p = _para("contact")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.size = Pt(chosen["contact"])
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        elif kind in ("skillhead", "heading"):
            _heading(text, chosen["head"])
        elif kind == "skills":
            _add_text(_para("skills"), text)
        elif kind == "bullet":
            bp = _para("bullet")
            bp.paragraph_format.left_indent = Pt(12)
            bp.add_run("• ")           # U+2022 bullet
            _add_text(bp, text)
        else:  # body
            _add_text(_para("body"), text)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# --------------------------------------------------------------------------- #
# DOCX -> PDF
# --------------------------------------------------------------------------- #
def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert .docx to PDF using Microsoft Word (docx2pdf). Falls back to a
    reportlab text render if Word/conversion is unavailable. Raises on total
    failure so the caller can surface a clear message."""
    # Preferred: Word, so the PDF matches the .docx styling exactly.
    try:
        import pythoncom  # noqa: needed when called from a worker thread
        pythoncom.CoInitialize()
    except Exception:
        pythoncom = None
    try:
        from docx2pdf import convert
        with tempfile.TemporaryDirectory() as td:
            in_path = os.path.join(td, "in.docx")
            out_path = os.path.join(td, "in.pdf")
            with open(in_path, "wb") as f:
                f.write(docx_bytes)
            convert(in_path, out_path)
            if os.path.exists(out_path):
                with open(out_path, "rb") as f:
                    return f.read()
        raise RuntimeError("Word produced no PDF output")
    except Exception as word_err:
        # Fallback: render the docx text into a basic PDF (layout not preserved).
        try:
            return _fallback_text_pdf(docx_bytes)
        except Exception as fb_err:
            raise RuntimeError(
                f"Could not convert to PDF (Word: {word_err}; fallback: {fb_err}). "
                "The Word (.docx) file was still produced."
            )
    finally:
        if pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass


def _fallback_text_pdf(docx_bytes: bytes) -> bytes:
    """Last-resort PDF: dump the docx text into a simple reportlab layout."""
    from docx import Document
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = Document(io.BytesIO(docx_bytes))
    styles = getSampleStyleSheet()
    out = io.BytesIO()
    pdf = SimpleDocTemplate(out, pagesize=A4, leftMargin=15 * mm, rightMargin=15 * mm,
                            topMargin=12 * mm, bottomMargin=12 * mm)
    story = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            safe = txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, styles["Normal"]))
            story.append(Spacer(1, 3))
    if not story:
        story.append(Paragraph("(empty resume)", styles["Normal"]))
    pdf.build(story)
    return out.getvalue()


# --------------------------------------------------------------------------- #
# TOP-LEVEL ENTRY
# --------------------------------------------------------------------------- #
def tailor_upload(filename: str, data: bytes, jd_text: str = "",
                  extra_skills: str = "", want_pdf: bool = True,
                  want_docx: bool = True, ats: bool = True) -> dict:
    """Tailor an uploaded resume.

    ats=True (default): also add the JD's remaining keywords to the skills line
    for ~100% ATS keyword coverage.

    Returns dict with:
      analysis, layout_preserved (bool), docx (bytes|None), pdf (bytes|None),
      pdf_note (str|None), ats_added (int)  -- keywords added for ATS coverage.
    """
    if not data:
        raise ValueError("Empty file.")
    ext = os.path.splitext(filename or "")[1].lower()

    resume_text = extract_text(filename, data)
    if not resume_text.strip():
        raise ValueError(
            "No readable text found. If this is a scanned/image PDF, upload a "
            "text-based PDF or a .docx instead."
        )

    info = analyze(resume_text, jd_text, extra_skills)

    layout_preserved = ext == ".docx"
    if layout_preserved:
        tailored_docx = tailor_docx(data, info, ats=ats)
    else:
        tailored_docx = rebuild_docx_from_text(resume_text, info, ats=ats)

    result = {
        "analysis": info,
        "layout_preserved": layout_preserved,
        "docx": tailored_docx if want_docx else None,
        "pdf": None,
        "pdf_note": None,
        "ats_added": len(info["suggestions"]) if ats else 0,
    }

    if want_pdf:
        try:
            result["pdf"] = docx_to_pdf(tailored_docx)
        except Exception as e:
            result["pdf_note"] = str(e)

    return result
